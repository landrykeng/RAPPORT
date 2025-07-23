import streamlit as st
import pandas as pd
import numpy as np
import os
import io
from datetime import datetime
import zipfile
import base64

# Imports pour les statistiques
from scipy import stats
from scipy.stats import chi2_contingency, fisher_exact, shapiro, normaltest
from scipy.stats import mannwhitneyu, kruskal, ttest_ind, f_oneway

# Imports pour les graphiques
from streamlit_echarts import st_echarts
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import plotly.io as pio

# Imports pour la génération de documents Word
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn

# Configuration de la page
st.set_page_config(
    page_title="📊 Analyseur Statistique Pro",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

# CSS personnalisé
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        color: #1f4e79;
        text-align: center;
        margin-bottom: 2rem;
        font-weight: bold;
    }
    .sub-header {
        font-size: 1.5rem;
        color: #0066cc;
        margin: 1rem 0;
    }
    .metric-container {
        background-color: #f0f8ff;
        padding: 1rem;
        border-radius: 10px;
        border-left: 5px solid #0066cc;
    }
    .success-box {
        background-color: #d4edda;
        padding: 1rem;
        border-radius: 5px;
        border: 1px solid #c3e6cb;
        color: #155724;
    }
    .warning-box {
        background-color: #fff3cd;
        padding: 1rem;
        border-radius: 5px;
        border: 1px solid #ffeaa7;
        color: #856404;
    }
    .info-box {
        background-color: #cce7ff;
        padding: 1rem;
        border-radius: 5px;
        border: 1px solid #99d6ff;
        color: #004080;
    }
    .analysis-container {
        border: 2px solid #e0e0e0;
        border-radius: 10px;
        padding: 1rem;
        margin: 1rem 0;
        background-color: #fafafa;
    }
    .palette-preview {
        display: flex;
        gap: 5px;
        margin: 5px 0;
    }
    .color-box {
        width: 20px;
        height: 20px;
        border-radius: 3px;
        border: 1px solid #ccc;
    }
</style>
""", unsafe_allow_html=True)

# Définition des palettes de couleurs
COLOR_PALETTES = {
    "Défaut (Verte/Orange)": [
        '#0C752C', '#F15A02', '#DBB300', '#FFEA2B', '#96CEB4', 
        '#FFEAA7', '#DDA0DD', '#98D8C8', '#F7DC6F', '#BB8FCE',
        '#85C1E9', '#F8C471', '#82E0AA', '#F1948A', '#F5B7B1'
    ],
    "Bleus océan": [
        '#1f77b4', '#aec7e8', '#ff7f0e', '#ffbb78', '#2ca02c',
        '#98df8a', '#d62728', '#ff9896', '#9467bd', '#c5b0d5',
        '#8c564b', '#c49c94', '#e377c2', '#f7b6d3', '#17becf'
    ],
    "Tons chauds": [
        '#d62728', '#ff7f0e', '#ff9896', '#ffbb78', '#e377c2',
        '#f7b6d3', '#bcbd22', '#dbdb8d', '#8c564b', '#c49c94',
        '#ff7f0e', '#ffbb78', '#2ca02c', '#98df8a', '#1f77b4'
    ],
    "Pastels doux": [
        '#FFB6C1', '#87CEEB', '#98FB98', '#F0E68C', '#DDA0DD',
        '#FFEAA7', '#96CEB4', '#F8C471', '#85C1E9', '#82E0AA',
        '#F1948A', '#BB8FCE', '#F7DC6F', '#98D8C8', '#F5B7B1'
    ],
    "Tons terreux": [
        '#8B4513', '#CD853F', '#D2691E', '#F4A460', '#DEB887',
        '#BC8F8F', '#F5DEB3', '#D3D3D3', '#A0522D', '#800000',
        '#B22222', '#DC143C', '#FF6347', '#FF4500', '#FF8C00'
    ],
    "Monochrome": [
        '#2F4F4F', '#696969', '#808080', '#A9A9A9', '#C0C0C0',
        '#D3D3D3', '#DCDCDC', '#F5F5F5', '#000000', '#1C1C1C',
        '#363636', '#4F4F4F', '#6B6B6B', '#878787', '#A3A3A3'
    ],
    "Arc-en-ciel": [
        '#FF0000', '#FF7F00', '#FFFF00', '#00FF00', '#0000FF',
        '#4B0082', '#9400D3', '#FF1493', '#00CED1', '#32CD32',
        '#FFD700', '#FF69B4', '#8A2BE2', '#00FA9A', '#FF6347'
    ]
}

# Configuration Plotly pour export d'images
pio.kaleido.scope.mathjax = None

# Fonction pour créer les dossiers de sortie
def create_output_folders():
    """Crée les dossiers de sortie s'ils n'existent pas"""
    folders = [
        "Sorties/Stat_Univariées",
        "Sorties/Stat_Bivariées",
        "Sorties/Exports"
    ]
    for folder in folders:
        if not os.path.exists(folder):
            os.makedirs(folder)

# Fonction pour afficher un aperçu de palette
def show_palette_preview(palette_name, colors):
    """Affiche un aperçu de la palette de couleurs"""
    color_boxes = ""
    for color in colors[:10]:
        color_boxes += f'<div class="color-box" style="background-color: {color};"></div>'
    
    st.markdown(f"""
    <div style="display: flex; align-items: center; gap: 10px;">
        <strong>{palette_name}:</strong>
        <div class="palette-preview">{color_boxes}</div>
    </div>
    """, unsafe_allow_html=True)

# Fonction pour lister les feuilles Excel
@st.cache_data
def get_excel_sheets(uploaded_file):
    """Récupère la liste des feuilles d'un fichier Excel"""
    try:
        excel_file = pd.ExcelFile(uploaded_file)
        return excel_file.sheet_names
    except Exception as e:
        st.error(f"Erreur lors de la lecture du fichier Excel: {str(e)}")
        return []

# Fonction de chargement des données
@st.cache_data
def load_data(uploaded_file, sheet_name=None):
    """Charge les données depuis un fichier uploadé"""
    try:
        if uploaded_file.name.endswith('.csv'):
            # Détection automatique du séparateur
            sample = str(uploaded_file.read(1024), 'utf-8')
            uploaded_file.seek(0)
            
            # Essayer différents séparateurs
            separators = [',', ';', '\t', '|']
            best_sep = ','
            max_cols = 0
            
            for sep in separators:
                try:
                    df_test = pd.read_csv(io.StringIO(sample), sep=sep, nrows=5)
                    if df_test.shape[1] > max_cols:
                        max_cols = df_test.shape[1]
                        best_sep = sep
                except:
                    continue
            
            df = pd.read_csv(uploaded_file, sep=best_sep)
            
        elif uploaded_file.name.endswith(('.xlsx', '.xls')):
            if sheet_name:
                df = pd.read_excel(uploaded_file, sheet_name=sheet_name)
            else:
                df = pd.read_excel(uploaded_file)
        else:
            st.error("Format de fichier non supporté. Utilisez CSV ou Excel.")
            return None
            
        return df
    except Exception as e:
        st.error(f"Erreur lors du chargement: {str(e)}")
        return None

# Fonction pour l'analyse descriptive univariée
def analyze_numeric_variables(df, numeric_vars, ponderation=None):
    """Analyse des variables numériques"""
    stats_dict = {}
    
    for var in numeric_vars:
        if ponderation is None:
            series = df[var].dropna()
            stats_dict[var] = {
                'Nombre': len(series),
                'Manquantes': df[var].isnull().sum(),
                '% Manquantes': round((df[var].isnull().sum() / len(df)) * 100, 2),
                'Moyenne': round(series.mean(), 3),
                'Médiane': round(series.median(), 3),
                'Écart-type': round(series.std(), 3),
                'Minimum': round(series.min(), 3),
                'Q1': round(series.quantile(0.25), 3),
                'Q3': round(series.quantile(0.75), 3),
                'Maximum': round(series.max(), 3),
                'Asymétrie': round(series.skew(), 4),
                'Aplatissement': round(series.kurtosis(), 4),
                'Total': round(series.sum(), 3)
            }
        else:
            # Analyse pondérée
            mask_valid = df[var].notna() & df[ponderation].notna()
            if mask_valid.sum() > 0:
                values = df.loc[mask_valid, var]
                weights = df.loc[mask_valid, ponderation]
                weighted_mean = np.average(values, weights=weights)
                weighted_var = np.average((values - weighted_mean)**2, weights=weights)
                
                stats_dict[var] = {
                    'Nombre': len(values),
                    'Effectif pondéré': round(weights.sum(), 1),
                    'Manquantes': df[var].isnull().sum(),
                    '% Manquantes': round((df[var].isnull().sum() / len(df)) * 100, 2),
                    'Moyenne': round(weighted_mean, 3),
                    'Écart-type': round(np.sqrt(weighted_var), 3),
                    'Minimum': round(values.min(), 3),
                    'Maximum': round(values.max(), 3),
                    'Total pondéré': round(np.sum(values * weights), 3)
                }
    
    return pd.DataFrame(stats_dict).T

# Fonctions de graphiques avec streamlit-echarts ET export Plotly

def create_pie_chart_echarts(value_counts, var_name, save_folder, colors, index=0):
    """Crée un graphique en secteurs avec streamlit-echarts et l'exporte en PNG
    Alternance entre disque (index pair) et anneau (index impair)"""
    
    # Alternance entre disque et anneau selon l'index
    hole_size = 0 if index % 2 == 0 else 0.4
    chart_type = "disque" if hole_size == 0 else "anneau"
    
    # Préparation des données
    data = [{"value": int(val), "name": str(name)} for name, val in value_counts.items()]
    
    # Configuration ECharts
    option = {
        "tooltip": {"trigger": "item"},
        "legend": {
            "orient": "vertical",
            "left": "left",
            "textStyle": {"fontSize": 12}
        },
        "series": [
            {
                "name": var_name,
                "type": "pie",
                "radius": ["40%" if hole_size > 0 else "0%", "70%"],
                "center": ["60%", "50%"],
                "data": data,
                "emphasis": {
                    "itemStyle": {
                        "shadowBlur": 10,
                        "shadowOffsetX": 0,
                        "shadowColor": "rgba(0, 0, 0, 0.5)"
                    }
                },
                "label": {
                    "show": True,
                    "formatter": "{b}: {c} ({d}%)",
                    "fontSize": 11
                },
                "labelLine": {"show": True},
                "color": colors[:len(data)]
            }
        ],
        "title": {
            "text": f"📊 {var_name}",
            "left": "center",
            "textStyle": {"fontSize": 16, "color": "#2E4057"}
        }
    }
    
    # Affichage avec streamlit-echarts
    st_echarts(options=option, height="500px", key=f"pie_{var_name}_{index}")
    
    # Export en PNG avec Plotly
    if save_folder:
        export_pie_chart_plotly(value_counts, var_name, save_folder, colors, hole_size, chart_type)

def export_pie_chart_plotly(value_counts, var_name, save_folder, colors, hole_size, chart_type):
    """Exporte le graphique en secteurs en PNG avec Plotly"""
    fig = go.Figure(data=[go.Pie(
        labels=value_counts.index,
        values=value_counts.values,
        hole=hole_size,
        textinfo='label+percent+value',
        textfont_size=12,
        marker=dict(
            colors=colors[:len(value_counts)],
            line=dict(color='white', width=2)
        )
    )])
    
    fig.update_layout(
        title={
            'text': f"📊 {var_name}",
            'x': 0.5,
            'xanchor': 'center',
            'font': {'size': 16, 'color': '#2E4057'}
        },
        font=dict(family="Arial, sans-serif", size=12),
        showlegend=True,
        legend=dict(orientation="v", yanchor="middle", y=0.5, xanchor="left", x=0.95),
        margin=dict(l=50, r=150, t=80, b=50),
        paper_bgcolor='rgba(0,0,0,0)',
        plot_bgcolor='rgba(0,0,0,0)',
        height=500,
        width=800
    )
    
    # Export
    clean_name = "".join(c for c in var_name if c.isalnum() or c in (' ', '-', '_')).rstrip()
    clean_name = clean_name.replace(' ', '_')
    save_path = f"{save_folder}/pie_{chart_type}_{clean_name}.png"
    
    try:
        pio.write_image(fig, save_path, width=800, height=500)
    except Exception as e:
        st.warning(f"Impossible de sauvegarder le graphique: {str(e)}")

def create_bar_chart_echarts(value_counts, var_name, save_folder, colors, index=0):
    """Crée un graphique en barres avec streamlit-echarts et l'exporte en PNG
    Alternance entre vertical (index pair) et horizontal (index impair)"""
    
    # Limitation à 15 modalités pour la lisibilité
    if len(value_counts) > 15:
        value_counts_display = value_counts.head(15)
        title_suffix = f" (Top 15 sur {len(value_counts)} modalités)"
    else:
        value_counts_display = value_counts
        title_suffix = ""
    
    categories = [str(x) for x in value_counts_display.index]
    values = value_counts_display.values.tolist()
    percentages = (value_counts_display / value_counts_display.sum() * 100).round(1).tolist()
    
    # Alternance entre vertical et horizontal selon l'index
    is_horizontal = index % 2 == 1
    chart_orientation = "horizontal" if is_horizontal else "vertical"
    
    # Configuration ECharts selon l'orientation
    if is_horizontal:
        # Barres horizontales
        option = {
            "tooltip": {
                "trigger": "axis",
                "formatter": "{b}<br/>Effectif: {c}<br/>Pourcentage: {customdata}%"
            },
            "xAxis": {
                "type": "value",
                "axisLabel": {"fontSize": 11}
            },
            "yAxis": {
                "type": "category",
                "data": categories,
                "axisLabel": {"fontSize": 11}
            },
            "series": [
                {
                    "data": [{"value": val, "customdata": pct} for val, pct in zip(values, percentages)],
                    "type": "bar",
                    "itemStyle": {"color": colors[0]},
                    "label": {
                        "show": True,
                        "position": "right",
                        "fontSize": 11
                    }
                }
            ],
            "title": {
                "text": f"📊 {var_name}{title_suffix}",
                "left": "center",
                "textStyle": {"fontSize": 16, "color": "#2E4057"}
            },
            "grid": {
                "left": "20%",
                "right": "10%",
                "bottom": "10%",
                "top": "15%"
            }
        }
    else:
        # Barres verticales
        option = {
            "tooltip": {
                "trigger": "axis",
                "formatter": "{b}<br/>Effectif: {c}<br/>Pourcentage: {customdata}%"
            },
            "xAxis": {
                "type": "category",
                "data": categories,
                "axisLabel": {
                    "rotate": 45 if len(categories) > 8 else 0,
                    "fontSize": 11
                }
            },
            "yAxis": {
                "type": "value",
                "axisLabel": {"fontSize": 11}
            },
            "series": [
                {
                    "data": [{"value": val, "customdata": pct} for val, pct in zip(values, percentages)],
                    "type": "bar",
                    "itemStyle": {"color": colors[0]},
                    "label": {
                        "show": True,
                        "position": "top",
                        "fontSize": 11
                    }
                }
            ],
            "title": {
                "text": f"📊 {var_name}{title_suffix}",
                "left": "center",
                "textStyle": {"fontSize": 16, "color": "#2E4057"}
            },
            "grid": {
                "left": "10%",
                "right": "10%",
                "bottom": "20%",
                "top": "15%"
            }
        }
    
    # Affichage avec streamlit-echarts
    st_echarts(options=option, height="500px", key=f"bar_{var_name}_{index}")
    
    # Export en PNG avec Plotly
    if save_folder:
        export_bar_chart_plotly(value_counts_display, var_name, save_folder, colors, 
                               title_suffix, chart_orientation)

def export_bar_chart_plotly(value_counts, var_name, save_folder, colors, title_suffix, orientation):
    """Exporte le graphique en barres en PNG avec Plotly"""
    
    if orientation == "horizontal":
        # Barres horizontales
        fig = go.Figure(data=[go.Bar(
            y=value_counts.index,
            x=value_counts.values,
            text=value_counts.values,
            textposition='outside',
            textfont=dict(size=12, color='#2E4057'),
            marker=dict(color=colors[0], line=dict(color='white', width=1)),
            orientation='h',
            hovertemplate='<b>%{y}</b><br>Effectif: %{x}<br>Pourcentage: %{customdata:.1f}%<br><extra></extra>',
            customdata=(value_counts / value_counts.sum() * 100)
        )])
        
        fig.update_layout(
            title=dict(text=f"📊 {var_name}{title_suffix}", x=0.5, xanchor='center', font=dict(size=16, color='#2E4057')),
            yaxis=dict(title=var_name, tickfont=dict(size=12)),
            xaxis=dict(title="Effectif", tickfont=dict(size=12), gridcolor='rgba(128,128,128,0.2)'),
            font=dict(family="Arial, sans-serif"),
            paper_bgcolor='rgba(0,0,0,0)',
            plot_bgcolor='rgba(0,0,0,0)',
            margin=dict(l=150, r=50, t=80, b=80),
            height=500,
            width=800
        )
    else:
        # Barres verticales
        fig = go.Figure(data=[go.Bar(
            x=value_counts.index,
            y=value_counts.values,
            text=value_counts.values,
            textposition='outside',
            textfont=dict(size=12, color='#2E4057'),
            marker=dict(color=colors[0], line=dict(color='white', width=1)),
            hovertemplate='<b>%{x}</b><br>Effectif: %{y}<br>Pourcentage: %{customdata:.1f}%<br><extra></extra>',
            customdata=(value_counts / value_counts.sum() * 100)
        )])
        
        fig.update_layout(
            title=dict(text=f"📊 {var_name}{title_suffix}", x=0.5, xanchor='center', font=dict(size=16, color='#2E4057')),
            xaxis=dict(title=var_name, tickfont=dict(size=12), tickangle=45 if len(value_counts) > 8 else 0),
            yaxis=dict(title="Effectif", tickfont=dict(size=12), gridcolor='rgba(128,128,128,0.2)'),
            font=dict(family="Arial, sans-serif"),
            paper_bgcolor='rgba(0,0,0,0)',
            plot_bgcolor='rgba(0,0,0,0)',
            margin=dict(l=80, r=50, t=80, b=100),
            height=500,
            width=800
        )
    
    clean_name = "".join(c for c in var_name if c.isalnum() or c in (' ', '-', '_')).rstrip()
    clean_name = clean_name.replace(' ', '_')
    save_path = f"{save_folder}/bar_{orientation}_{clean_name}.png"
    
    try:
        pio.write_image(fig, save_path, width=800, height=500)
    except Exception as e:
        st.warning(f"Impossible de sauvegarder le graphique: {str(e)}")

def create_boxplot_echarts(df, var_cible, var_num, save_folder, colors):
    """Crée un boxplot avec streamlit-echarts et l'exporte en PNG"""
    
    groups = df.groupby(var_cible)[var_num].apply(list).to_dict()
    
    # Calculer les statistiques pour chaque groupe
    boxplot_data = []
    categories = []
    
    for group_name, group_values in groups.items():
        if len(group_values) > 0:
            q1 = np.percentile(group_values, 25)
            median = np.percentile(group_values, 50)
            q3 = np.percentile(group_values, 75)
            min_val = np.min(group_values)
            max_val = np.max(group_values)
            
            boxplot_data.append([min_val, q1, median, q3, max_val])
            categories.append(str(group_name))
    
    # Configuration ECharts
    option = {
        "title": {
            "text": f"📊 Distribution de {var_num} selon {var_cible}",
            "left": "center",
            "textStyle": {"fontSize": 16, "color": "#2E4057"}
        },
        "tooltip": {"trigger": "item"},
        "xAxis": {
            "type": "category",
            "data": categories,
            "axisLabel": {"fontSize": 11}
        },
        "yAxis": {
            "type": "value",
            "name": var_num,
            "axisLabel": {"fontSize": 11}
        },
        "series": [
            {
                "name": var_num,
                "type": "boxplot",
                "data": boxplot_data,
                "itemStyle": {"color": colors[0]}
            }
        ],
        "grid": {
            "left": "10%",
            "right": "10%",
            "bottom": "15%",
            "top": "15%"
        }
    }
    
    # Affichage avec streamlit-echarts
    st_echarts(options=option, height="500px", key=f"boxplot_{var_cible}_{var_num}")
    
    # Export en PNG avec Plotly
    if save_folder:
        export_boxplot_plotly(df, var_cible, var_num, save_folder, colors)

def export_boxplot_plotly(df, var_cible, var_num, save_folder, colors):
    """Exporte le boxplot en PNG avec Plotly"""
    fig = px.box(df, x=var_cible, y=var_num, color=var_cible, color_discrete_sequence=colors,
                 title=f"📊 Distribution de {var_num} selon {var_cible}")
    
    fig.update_layout(
        title={'x': 0.5, 'xanchor': 'center', 'font': {'size': 16, 'color': '#2E4057'}},
        xaxis_title=var_cible, yaxis_title=var_num, showlegend=False, height=500,
        font=dict(family="Arial, sans-serif"), paper_bgcolor='rgba(0,0,0,0)', plot_bgcolor='rgba(0,0,0,0)'
    )
    
    clean_cible = "".join(c for c in var_cible if c.isalnum() or c in (' ', '-', '_')).rstrip()
    clean_num = "".join(c for c in var_num if c.isalnum() or c in (' ', '-', '_')).rstrip()
    clean_cible = clean_cible.replace(' ', '_')
    clean_num = clean_num.replace(' ', '_')
    save_path = f"{save_folder}/boxplot_{clean_num}_par_{clean_cible}.png"
    
    try:
        pio.write_image(fig, save_path, width=800, height=500)
    except Exception as e:
        st.warning(f"Impossible de sauvegarder le graphique: {str(e)}")

def create_stacked_bar_echarts(df, var_cible, var_cat, save_folder, colors):
    """Crée un graphique en barres empilées avec streamlit-echarts et l'exporte en PNG"""
    
    # Calcul des pourcentages
    crosstab_pct = pd.crosstab(df[var_cat], df[var_cible], normalize='index') * 100
    
    categories = [str(x) for x in crosstab_pct.index]
    series_data = []
    
    for i, col in enumerate(crosstab_pct.columns):
        series_data.append({
            "name": str(col),
            "type": "bar",
            "stack": "total",
            "data": crosstab_pct[col].round(2).tolist(),
            "itemStyle": {"color": colors[i % len(colors)]},
            "label": {
                "show": True,
                "position": "inside",
                "formatter": "{c}%",
                "fontSize": 10
            }
        })
    
    # Configuration ECharts
    option = {
        "title": {
            "text": f"📊 Répartition de {var_cat} selon {var_cible} (%)",
            "left": "center",
            "textStyle": {"fontSize": 16, "color": "#2E4057"}
        },
        "tooltip": {
            "trigger": "axis",
            "formatter": "{b}<br/>{a}: {c}%"
        },
        "legend": {
            "data": [str(x) for x in crosstab_pct.columns],
            "top": "bottom",
            "textStyle": {"fontSize": 12}
        },
        "xAxis": {
            "type": "category",
            "data": categories,
            "axisLabel": {"fontSize": 11}
        },
        "yAxis": {
            "type": "value",
            "name": "Pourcentage (%)",
            "axisLabel": {"fontSize": 11}
        },
        "series": series_data,
        "grid": {
            "left": "10%",
            "right": "10%",
            "bottom": "20%",
            "top": "15%"
        }
    }
    
    # Affichage avec streamlit-echarts
    st_echarts(options=option, height="500px", key=f"stacked_{var_cat}_{var_cible}")
    
    # Export en PNG avec Plotly
    if save_folder:
        export_stacked_bar_plotly(df, var_cible, var_cat, save_folder, colors)

def export_stacked_bar_plotly(df, var_cible, var_cat, save_folder, colors):
    """Exporte le graphique en barres empilées en PNG avec Plotly"""
    crosstab_pct = pd.crosstab(df[var_cat], df[var_cible], normalize='index') * 100
    
    fig = go.Figure()
    
    for i, col in enumerate(crosstab_pct.columns):
        fig.add_trace(go.Bar(
            name=str(col), x=crosstab_pct.index, y=crosstab_pct[col],
            marker_color=colors[i % len(colors)],
            text=crosstab_pct[col].round(1).values, texttemplate='%{text}%', textposition='inside'
        ))
    
    fig.update_layout(
        title={'text': f"📊 Répartition de {var_cat} selon {var_cible} (%)", 'x': 0.5, 'xanchor': 'center',
               'font': {'size': 16, 'color': '#2E4057'}},
        xaxis_title=var_cat, yaxis_title="Pourcentage (%)", barmode='stack', height=500,
        font=dict(family="Arial, sans-serif", size=12), legend=dict(title=var_cible),
        paper_bgcolor='rgba(0,0,0,0)', plot_bgcolor='rgba(0,0,0,0)', margin=dict(l=80, r=50, t=80, b=80)
    )
    
    clean_cible = "".join(c for c in var_cible if c.isalnum() or c in (' ', '-', '_')).rstrip()
    clean_cat = "".join(c for c in var_cat if c.isalnum() or c in (' ', '-', '_')).rstrip()
    clean_cible = clean_cible.replace(' ', '_')
    clean_cat = clean_cat.replace(' ', '_')
    save_path = f"{save_folder}/stacked_{clean_cat}_selon_{clean_cible}.png"
    
    try:
        pio.write_image(fig, save_path, width=800, height=500)
    except Exception as e:
        st.warning(f"Impossible de sauvegarder le graphique: {str(e)}")

def create_heatmap_echarts(df, var_cible, var_cat, save_folder, colors):
    """Crée une heatmap avec streamlit-echarts et l'exporte en PNG"""
    
    crosstab_pct = pd.crosstab(df[var_cible], df[var_cat], normalize='index') * 100
    
    # Préparer les données pour la heatmap
    data = []
    for i, row_name in enumerate(crosstab_pct.index):
        for j, col_name in enumerate(crosstab_pct.columns):
            data.append([j, i, round(crosstab_pct.loc[row_name, col_name], 1)])
    
    # Configuration ECharts
    option = {
        "title": {
            "text": f"🔥 Heatmap: {var_cible} vs {var_cat} (%)",
            "left": "center",
            "textStyle": {"fontSize": 16, "color": "#2E4057"}
        },
        "tooltip": {
            "position": "top",
            "formatter": "{b}<br/>{a}: {c}%"
        },
        "grid": {"height": "60%", "top": "10%"},
        "xAxis": {
            "type": "category",
            "data": [str(x) for x in crosstab_pct.columns],
            "splitArea": {"show": True},
            "axisLabel": {"fontSize": 11}
        },
        "yAxis": {
            "type": "category",
            "data": [str(x) for x in crosstab_pct.index],
            "splitArea": {"show": True},
            "axisLabel": {"fontSize": 11}
        },
        "visualMap": {
            "min": 0,
            "max": crosstab_pct.values.max(),
            "calculable": True,
            "orient": "horizontal",
            "left": "center",
            "bottom": "15%",
            "inRange": {"color": ["#f7f7f7", colors[0]]}
        },
        "series": [
            {
                "name": "Pourcentage",
                "type": "heatmap",
                "data": data,
                "label": {"show": True, "fontSize": 11}
            }
        ]
    }
    
    # Affichage avec streamlit-echarts
    st_echarts(options=option, height="500px", key=f"heatmap_{var_cible}_{var_cat}")
    
    # Export en PNG avec Plotly
    if save_folder:
        export_heatmap_plotly(df, var_cible, var_cat, save_folder)

def export_heatmap_plotly(df, var_cible, var_cat, save_folder):
    """Exporte la heatmap en PNG avec Plotly"""
    crosstab_pct = pd.crosstab(df[var_cible], df[var_cat], normalize='index') * 100
    
    fig = go.Figure(data=go.Heatmap(
        z=crosstab_pct.values, x=crosstab_pct.columns, y=crosstab_pct.index,
        colorscale='Blues', text=crosstab_pct.round(1).values, texttemplate='%{text}%',
        textfont={"size": 12}, hoverongaps=False, colorbar=dict(title="Pourcentage (%)")
    ))
    
    fig.update_layout(
        title={'text': f"🔥 Heatmap: {var_cible} vs {var_cat} (%)", 'x': 0.5, 'xanchor': 'center',
               'font': {'size': 16, 'color': '#2E4057'}},
        xaxis_title=var_cat, yaxis_title=var_cible, height=500, font=dict(family="Arial, sans-serif"),
        paper_bgcolor='rgba(0,0,0,0)', plot_bgcolor='rgba(0,0,0,0)'
    )
    
    clean_cible = "".join(c for c in var_cible if c.isalnum() or c in (' ', '-', '_')).rstrip()
    clean_cat = "".join(c for c in var_cat if c.isalnum() or c in (' ', '-', '_')).rstrip()
    clean_cible = clean_cible.replace(' ', '_')
    clean_cat = clean_cat.replace(' ', '_')
    save_path = f"{save_folder}/heatmap_{clean_cible}_vs_{clean_cat}.png"
    
    try:
        pio.write_image(fig, save_path, width=800, height=500)
    except Exception as e:
        st.warning(f"Impossible de sauvegarder le graphique: {str(e)}")

# Fonctions pour la génération de documents Word
def create_word_document_header(doc, title, subtitle=None):
    """Crée l'en-tête du document Word avec mise en forme"""
    # Titre principal
    title_para = doc.add_heading(title, 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if subtitle:
        subtitle_para = doc.add_paragraph(subtitle)
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle_para.runs[0]
        subtitle_run.font.size = Pt(14)
        subtitle_run.font.italic = True
    
    # Ligne de séparation
    doc.add_paragraph("=" * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

def add_dataframe_to_doc(doc, df, title=None):
    """Ajoute un DataFrame au document Word sous forme de tableau"""
    if title:
        doc.add_heading(title, level=3)
    
    # Créer le tableau
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # En-têtes
    header_cells = table.rows[0].cells
    for i, column in enumerate(df.columns):
        header_cells[i].text = str(column)
        header_cells[i].paragraphs[0].runs[0].font.bold = True
    
    # Données
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)
    
    doc.add_paragraph()

def add_image_to_doc(doc, image_path, title=None, width=Inches(6)):
    """Ajoute une image au document Word"""
    if title:
        doc.add_heading(title, level=3)
    
    try:
        if os.path.exists(image_path):
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            run.add_picture(image_path, width=width)
            doc.add_paragraph()
        else:
            doc.add_paragraph(f"⚠️ Image non trouvée: {image_path}")
    except Exception as e:
        doc.add_paragraph(f"⚠️ Erreur lors de l'insertion de l'image: {str(e)}")

def create_word_report(df, numeric_vars, categorical_vars, resultats_bivariees=None, 
                      var_cible=None, ponderation=None, palette_name=None, selected_sheet=None):
    """Crée un rapport Word complet avec toutes les analyses"""
    doc = Document()
    
    # En-tête du document
    create_word_document_header(
        doc, 
        "📊 RAPPORT D'ANALYSE STATISTIQUE",
        f"Généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}"
    )
    
    # Informations générales
    doc.add_heading("📋 Informations générales", level=1)
    
    info_data = [
        ["Nombre de lignes", str(df.shape[0])],
        ["Nombre de colonnes", str(df.shape[1])],
        ["Variables numériques", str(len(numeric_vars))],
        ["Variables qualitatives", str(len(categorical_vars))],
        ["Feuille Excel", selected_sheet if selected_sheet else "N/A (fichier CSV)"],
        ["Pondération utilisée", ponderation if ponderation else "Aucune"],
        ["Palette de couleurs", palette_name if palette_name else "Défaut"],
        ["Date d'analyse", datetime.now().strftime('%d/%m/%Y à %H:%M')]
    ]
    
    info_table = doc.add_table(rows=1, cols=2)
    info_table.style = 'Light List Accent 1'
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # En-têtes du tableau d'informations
    header_cells = info_table.rows[0].cells
    header_cells[0].text = "Paramètre"
    header_cells[1].text = "Valeur"
    for cell in header_cells:
        cell.paragraphs[0].runs[0].font.bold = True
    
    # Données du tableau d'informations
    for param, valeur in info_data:
        row_cells = info_table.add_row().cells
        row_cells[0].text = param
        row_cells[1].text = valeur
    
    doc.add_page_break()
    
    # Analyses univariées
    doc.add_heading("📊 ANALYSES UNIVARIÉES", level=1)
    
    # Variables numériques
    if numeric_vars:
        doc.add_heading("🔢 Variables numériques", level=2)
        
        # Calculer les statistiques
        stats_table = analyze_numeric_variables(df, numeric_vars, ponderation)
        add_dataframe_to_doc(doc, stats_table, "Statistiques descriptives")
        
        doc.add_paragraph("Les statistiques ci-dessus présentent les mesures de tendance centrale, "
                         "de dispersion et de forme pour chaque variable numérique.")
        if ponderation:
            doc.add_paragraph(f"⚖️ Analyses pondérées avec la variable: {ponderation}")
    
    # Variables qualitatives
    if categorical_vars:
        doc.add_heading("📝 Variables qualitatives", level=2)
        
        for i, var in enumerate(categorical_vars):
            doc.add_heading(f"Variable: {var}", level=3)
            
            # Calculer les fréquences
            if ponderation is None:
                value_counts = df[var].value_counts()
                total = value_counts.sum()
            else:
                mask_valid = df[var].notna() & df[ponderation].notna()
                if mask_valid.sum() > 0:
                    value_counts = df[mask_valid].groupby(var)[ponderation].sum().sort_values(ascending=False)
                    total = value_counts.sum()
                else:
                    continue
            
            # Tableau de fréquences
            freq_table = pd.DataFrame({
                'Modalité': value_counts.index,
                'Effectif': value_counts.values,
                'Fréquence (%)': (value_counts / total * 100).round(2).values
            })
            
            add_dataframe_to_doc(doc, freq_table, f"Répartition de {var}")
            
            # Informations sur la variable
            n_categories = len(value_counts)
            missing_count = df[var].isnull().sum()
            
            info_text = f"• Nombre de modalités: {n_categories}\n"
            info_text += f"• Valeurs manquantes: {missing_count} ({(missing_count/len(df)*100):.1f}%)\n"
            if ponderation:
                info_text += f"• Effectif pondéré total: {total:.0f}\n"
            
            doc.add_paragraph(info_text)
            
            # Ajouter le graphique si disponible
            clean_name = "".join(c for c in var if c.isalnum() or c in (' ', '-', '_')).rstrip()
            clean_name = clean_name.replace(' ', '_')
            
            if n_categories <= 3:
                # Graphique en secteurs (alternance disque/anneau)
                chart_type = "disque" if i % 2 == 0 else "anneau"
                image_path = f"Sorties/Stat_Univariées/pie_{chart_type}_{clean_name}.png"
                add_image_to_doc(doc, image_path, f"Graphique: {var}")
            else:
                # Graphique en barres (alternance vertical/horizontal)
                chart_orientation = "vertical" if i % 2 == 0 else "horizontal"
                image_path = f"Sorties/Stat_Univariées/bar_{chart_orientation}_{clean_name}.png"
                add_image_to_doc(doc, image_path, f"Graphique: {var}")
    
    # Analyses bivariées
    if resultats_bivariees and var_cible:
        doc.add_page_break()
        doc.add_heading("🔗 ANALYSES BIVARIÉES", level=1)
        
        doc.add_paragraph(f"Variable cible: {var_cible}")
        doc.add_paragraph(f"Nombre d'analyses: {len(resultats_bivariees)}")
        doc.add_paragraph()
        
        for var_name, resultats in resultats_bivariees.items():
            doc.add_heading(f"Analyse: {var_cible} × {var_name}", level=2)
            
            # Type d'analyse et informations
            doc.add_paragraph(f"Type d'analyse: {resultats['type_analyse']}")
            doc.add_paragraph(f"Nombre d'observations: {resultats['n_observations']}")
            
            # Tableau principal
            if 'tableau_principal' in resultats and resultats['tableau_principal'] is not None:
                add_dataframe_to_doc(doc, resultats['tableau_principal'], "Tableau principal")
            
            # Tableau secondaire
            if 'tableau_secondaire' in resultats and resultats['tableau_secondaire'] is not None:
                add_dataframe_to_doc(doc, resultats['tableau_secondaire'], "Tableau des pourcentages")
            
            # Tests statistiques
            if 'tests' in resultats and resultats['tests']:
                doc.add_heading("Tests statistiques", level=3)
                tests_df = pd.DataFrame(resultats['tests'])
                add_dataframe_to_doc(doc, tests_df)
                
                # Résumé du test principal
                if resultats.get('test_principal'):
                    doc.add_paragraph(f"Test principal: {resultats['test_principal']}")
                    if resultats.get('p_value_principal'):
                        doc.add_paragraph(f"P-value: {resultats['p_value_principal']}")
                        if resultats.get('significatif_principal'):
                            doc.add_paragraph("✅ Résultat significatif (p < 0.05)")
                        else:
                            doc.add_paragraph("❌ Résultat non significatif (p ≥ 0.05)")
            
            # Ajouter les graphiques
            clean_cible = "".join(c for c in var_cible if c.isalnum() or c in (' ', '-', '_')).rstrip()
            clean_var = "".join(c for c in var_name if c.isalnum() or c in (' ', '-', '_')).rstrip()
            clean_cible = clean_cible.replace(' ', '_')
            clean_var = clean_var.replace(' ', '_')
            
            if resultats['type_analyse'] == 'Numérique vs Catégorielle':
                # Boxplot
                image_path = f"Sorties/Stat_Bivariées/boxplot_{clean_var}_par_{clean_cible}.png"
                add_image_to_doc(doc, image_path, "Distribution par groupe (Boxplot)", width=Inches(5))
                
            elif resultats['type_analyse'] == 'Catégorielle vs Catégorielle':
                # Barres empilées
                image_path = f"Sorties/Stat_Bivariées/stacked_{clean_var}_selon_{clean_cible}.png"
                add_image_to_doc(doc, image_path, "Répartition (Barres empilées)", width=Inches(5))
                
                # Heatmap
                image_path = f"Sorties/Stat_Bivariées/heatmap_{clean_cible}_vs_{clean_var}.png"
                add_image_to_doc(doc, image_path, "Carte de chaleur (Heatmap)", width=Inches(5))
                
            elif resultats['type_analyse'] == 'Numérique vs Numérique':
                # Nuage de points
                image_path = f"Sorties/Stat_Bivariées/scatter_{clean_cible}_vs_{clean_var}.png"
                add_image_to_doc(doc, image_path, "Nuage de points", width=Inches(5))
            
            doc.add_paragraph()
    
    # Conclusion
    doc.add_page_break()
    doc.add_heading("📋 CONCLUSION", level=1)
    
    conclusion_text = "Ce rapport présente une analyse statistique complète des données fournies. "
    conclusion_text += f"L'analyse porte sur {df.shape[0]} observations et {df.shape[1]} variables, "
    conclusion_text += f"dont {len(numeric_vars)} variables numériques et {len(categorical_vars)} variables qualitatives.\n\n"
    
    if ponderation:
        conclusion_text += f"Les analyses ont été pondérées en utilisant la variable '{ponderation}'.\n\n"
    
    if resultats_bivariees:
        conclusion_text += f"Les analyses bivariées ont permis d'étudier les relations entre la variable cible '{var_cible}' "
        conclusion_text += f"et {len(resultats_bivariees)} autres variables.\n\n"
    
    conclusion_text += "Les graphiques et tableaux présentés dans ce rapport permettent une compréhension "
    conclusion_text += "approfondie de la structure et des relations présentes dans les données."
    
    doc.add_paragraph(conclusion_text)
    
    # Note technique
    doc.add_heading("📌 Note technique", level=2)
    tech_note = "Ce rapport a été généré automatiquement par l'Analyseur Statistique Pro. "
    tech_note += "Les analyses statistiques ont été réalisées avec Python et les bibliothèques "
    tech_note += "pandas, scipy et numpy. Les graphiques ont été créés avec plotly et streamlit-echarts."
    
    doc.add_paragraph(tech_note)
    
    return doc

# Fonction pour créer le rapport Excel des analyses bivariées
def create_bivariate_excel_report(resultats_analyses, var_cible, ponderation, palette_name, selected_sheet=None):
    """Crée un rapport Excel complet des analyses bivariées"""
    excel_buffer = io.BytesIO()
    
    with pd.ExcelWriter(excel_buffer, engine='openpyxl') as writer:
        
        # Feuille de synthèse
        synthese_data = []
        for var_name, resultats in resultats_analyses.items():
            synthese_data.append({
                'Variable': var_name,
                'Type d\'analyse': resultats['type_analyse'],
                'Nombre d\'observations': resultats['n_observations'],
                'Test principal': resultats.get('test_principal', 'N/A'),
                'P-value': resultats.get('p_value_principal', 'N/A'),
                'Significatif': resultats.get('significatif_principal', 'N/A')
            })
        
        synthese_df = pd.DataFrame(synthese_data)
        synthese_df.to_excel(writer, sheet_name='00_Synthese', index=False)
        
        # Informations générales
        info_df = pd.DataFrame({
            'Paramètre': ['Variable cible', 'Feuille Excel', 'Pondération', 'Palette de couleurs', 'Date d\'analyse', 'Nombre d\'analyses'],
            'Valeur': [var_cible, selected_sheet if selected_sheet else 'N/A (CSV)', 
                      ponderation if ponderation else 'Aucune', palette_name, 
                      datetime.now().strftime('%Y-%m-%d %H:%M'), len(resultats_analyses)]
        })
        info_df.to_excel(writer, sheet_name='00_Informations', index=False)
        
        # Feuille pour chaque analyse
        for i, (var_name, resultats) in enumerate(resultats_analyses.items(), 1):
            sheet_name = f"{i:02d}_{var_name[:25]}"  # Limiter la longueur du nom
            
            # Écrire les différents tableaux sur la même feuille
            current_row = 0
            
            # Titre de l'analyse
            title_df = pd.DataFrame({
                'Analyse': [f"{var_cible} × {var_name}"],
                'Type': [resultats['type_analyse']],
                'Observations': [resultats['n_observations']]
            })
            title_df.to_excel(writer, sheet_name=sheet_name, startrow=current_row, index=False)
            current_row += 3
            
            # Tableau principal (crosstab ou stats descriptives)
            if 'tableau_principal' in resultats:
                # Ajouter un en-tête
                header_df = pd.DataFrame({'': ['TABLEAU PRINCIPAL']})
                header_df.to_excel(writer, sheet_name=sheet_name, startrow=current_row, index=False, header=False)
                current_row += 1
                
                resultats['tableau_principal'].to_excel(writer, sheet_name=sheet_name, startrow=current_row)
                current_row += len(resultats['tableau_principal']) + 3
            
            # Tests statistiques
            if 'tests' in resultats and resultats['tests']:
                header_df = pd.DataFrame({'': ['TESTS STATISTIQUES']})
                header_df.to_excel(writer, sheet_name=sheet_name, startrow=current_row, index=False, header=False)
                current_row += 1
                
                tests_df = pd.DataFrame(resultats['tests'])
                tests_df.to_excel(writer, sheet_name=sheet_name, startrow=current_row, index=False)
                current_row += len(tests_df) + 3
            
            # Tableau supplémentaire (ex: pourcentages pour les variables catégorielles)
            if 'tableau_secondaire' in resultats:
                header_df = pd.DataFrame({'': ['TABLEAU SECONDAIRE']})
                header_df.to_excel(writer, sheet_name=sheet_name, startrow=current_row, index=False, header=False)
                current_row += 1
                
                resultats['tableau_secondaire'].to_excel(writer, sheet_name=sheet_name, startrow=current_row)
                current_row += len(resultats['tableau_secondaire']) + 3
    
    return excel_buffer

def analyze_numeric_vs_categorical(df_clean, var_cible, var_num, ponderation, save_folder, colors):
    """Analyse croisée variable catégorielle vs numérique avec affichage complet et retour de données"""
    
    # Dictionnaire pour stocker les résultats
    resultats = {
        'type_analyse': 'Numérique vs Catégorielle',
        'n_observations': len(df_clean),
        'tests': [],
        'tableau_principal': None,
        'test_principal': None,
        'p_value_principal': None,
        'significatif_principal': None
    }
    
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.write("📋 **Statistiques descriptives par groupe:**")
        
        # Tableau des statistiques descriptives par groupe
        if ponderation is None:
            tableau_desc = df_clean.groupby(var_cible)[var_num].agg([
                'count', 'mean', 'std', 'min', 'median', 'max'
            ]).round(3)
        else:
            # Version simplifiée pour la pondération
            stats_list = []
            for groupe_name in df_clean[var_cible].unique():
                groupe_data = df_clean[df_clean[var_cible] == groupe_name]
                values = groupe_data[var_num]
                weights = groupe_data[ponderation]
                
                if len(values) > 0:
                    weighted_mean = np.average(values, weights=weights)
                    stats_list.append({
                        'count': len(values),
                        'mean': weighted_mean,
                        'min': values.min(),
                        'median': np.median(values),  # Approximation
                        'max': values.max()
                    })
                else:
                    stats_list.append({
                        'count': 0, 'mean': np.nan, 'min': np.nan, 
                        'median': np.nan, 'max': np.nan
                    })
            
            tableau_desc = pd.DataFrame(stats_list, index=sorted(df_clean[var_cible].unique())).round(3)
        
        resultats['tableau_principal'] = tableau_desc
        st.dataframe(tableau_desc, use_container_width=True)
        
        # Tests statistiques
        st.write("🧪 **Tests statistiques:**")
        groupes = df_clean.groupby(var_cible)[var_num]
        groupes_data = [groupe_data.values for _, groupe_data in groupes if len(groupe_data) >= 3]
        
        if len(groupes_data) >= 2:
            try:
                if len(groupes_data) == 2:
                    stat, p_value = ttest_ind(groupes_data[0], groupes_data[1])
                    test_name = "Test t de Student"
                else:
                    stat, p_value = f_oneway(*groupes_data)
                    test_name = "ANOVA (Test F)"
                
                # Stockage du test principal
                resultats['test_principal'] = test_name
                resultats['p_value_principal'] = round(p_value, 4)
                resultats['significatif_principal'] = p_value < 0.05
                
                st.write(f"• **{test_name}:** {stat:.4f}")
                st.write(f"• **p-value:** {p_value:.4f}")
                st.write(f"• **Résultat:** {'Différence significative' if p_value < 0.05 else 'Pas de différence significative'}")
                
                # Ajouter aux tests
                resultats['tests'].append({
                    'Test': test_name,
                    'Statistique': round(stat, 4),
                    'p-value': round(p_value, 4),
                    'Significatif': p_value < 0.05,
                    'Interprétation': 'Différence significative' if p_value < 0.05 else 'Pas de différence significative'
                })
                
                # Test de normalité pour chaque groupe
                st.write("• **Tests de normalité:**")
                tests_normalite = []
                for groupe_name, groupe_data in groupes:
                    if len(groupe_data) >= 3:
                        if len(groupe_data) < 50:
                            stat_norm, p_norm = shapiro(groupe_data)
                            test_norm = "Shapiro-Wilk"
                        else:
                            stat_norm, p_norm = normaltest(groupe_data)
                            test_norm = "D'Agostino-Pearson"
                        
                        est_normal = p_norm > 0.05
                        st.write(f"  - {groupe_name}: {test_norm} p={p_norm:.3f} {'✅ Normal' if est_normal else '❌ Non-normal'}")
                        
                        tests_normalite.append({
                            'Test': f"{test_norm} ({groupe_name})",
                            'Statistique': round(stat_norm, 4),
                            'p-value': round(p_norm, 4),
                            'Significatif': not est_normal,
                            'Interprétation': 'Normal' if est_normal else 'Non-normal'
                        })
                
                resultats['tests'].extend(tests_normalite)
                        
            except Exception as e:
                st.write(f"• Erreur dans les tests: {str(e)}")
                resultats['tests'].append({
                    'Test': 'Erreur',
                    'Statistique': 'N/A',
                    'p-value': 'N/A',
                    'Significatif': False,
                    'Interprétation': str(e)
                })
        else:
            st.write("• Pas assez de groupes pour les tests")
            resultats['tests'].append({
                'Test': 'Impossible',
                'Statistique': 'N/A',
                'p-value': 'N/A',
                'Significatif': False,
                'Interprétation': 'Pas assez de groupes'
            })
    
    with col2:
        st.write("📊 **Graphiques de comparaison:**")
        
        # Affichage des graphiques dans la même colonne
        create_boxplot_echarts(df_clean, var_cible, var_num, save_folder, colors)
    
    return resultats

def analyze_categorical_vs_categorical(df_clean, var_cible, var_cat, ponderation, save_folder, colors):
    """Analyse croisée entre deux variables catégorielles avec affichage complet et retour de données"""
    
    # Dictionnaire pour stocker les résultats
    resultats = {
        'type_analyse': 'Catégorielle vs Catégorielle',
        'n_observations': len(df_clean),
        'tests': [],
        'tableau_principal': None,
        'tableau_secondaire': None,
        'test_principal': None,
        'p_value_principal': None,
        'significatif_principal': None
    }
    
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.write("📋 **Tableau croisé:**")
        
        # Tableau croisé
        if ponderation is None:
            tab_fi = pd.crosstab(df_clean[var_cible], df_clean[var_cat], margins=True)
            tab_pct = pd.crosstab(df_clean[var_cible], df_clean[var_cat], normalize='index') * 100
        else:
            tab_fi = pd.crosstab(
                df_clean[var_cible], df_clean[var_cat], 
                values=df_clean[ponderation], aggfunc='sum', margins=True
            ).fillna(0).round(1)
            tab_pct = pd.crosstab(
                df_clean[var_cible], df_clean[var_cat], 
                values=df_clean[ponderation], aggfunc='sum', normalize='index'
            ).fillna(0) * 100
        
        resultats['tableau_principal'] = tab_fi
        resultats['tableau_secondaire'] = tab_pct.round(2)
        
        st.dataframe(tab_fi, use_container_width=True)
        
        st.write("📊 **Pourcentages par ligne:**")
        st.dataframe(tab_pct.round(2), use_container_width=True)
        
        # Test du Chi2
        st.write("🧪 **Test d'indépendance (Chi2):**")
        try:
            if ponderation is None:
                observed = pd.crosstab(df_clean[var_cible], df_clean[var_cat])
            else:
                observed = pd.crosstab(
                    df_clean[var_cible], df_clean[var_cat], 
                    values=df_clean[ponderation], aggfunc='sum'
                ).fillna(0)
            
            chi2_stat, p_value, dof, expected = chi2_contingency(observed)
            
            # Stockage du test principal
            resultats['test_principal'] = "Test du Chi2"
            resultats['p_value_principal'] = round(p_value, 4)
            resultats['significatif_principal'] = p_value < 0.05
            
            st.write(f"• **Chi2:** {chi2_stat:.4f}")
            st.write(f"• **p-value:** {p_value:.4f}")
            st.write(f"• **Degrés de liberté:** {dof}")
            st.write(f"• **Résultat:** {'Association significative' if p_value < 0.05 else 'Association non significative'}")
            
            # Critères de validité
            expected_min = expected.min()
            cells_less_than_5 = (expected < 5).sum().sum()
            total_cells = expected.size
            percent_less_than_5 = (cells_less_than_5 / total_cells) * 100
            
            st.write("• **Critères de validité:**")
            st.write(f"  - Effectif théorique minimum: {expected_min:.2f}")
            st.write(f"  - Cellules < 5: {cells_less_than_5}/{total_cells} ({percent_less_than_5:.1f}%)")
            test_valide = expected_min >= 1 and percent_less_than_5 <= 20
            st.write(f"  - Test valide: {'✅ Oui' if test_valide else '❌ Non'}")
            
            # Ajouter le test principal
            resultats['tests'].append({
                'Test': 'Test du Chi2',
                'Statistique': round(chi2_stat, 4),
                'p-value': round(p_value, 4),
                'Degrés de liberté': dof,
                'Significatif': p_value < 0.05,
                'Interprétation': 'Association significative' if p_value < 0.05 else 'Association non significative'
            })
            
            # Ajouter les critères de validité
            resultats['tests'].append({
                'Test': 'Critères de validité',
                'Effectif min théorique': round(expected_min, 2),
                'Cellules < 5': f"{cells_less_than_5}/{total_cells}",
                'Pourcentage < 5': f"{percent_less_than_5:.1f}%",
                'Test valide': test_valide,
                'Interprétation': 'Test valide' if test_valide else 'Test non valide'
            })
            
            # Coefficient de Cramér's V
            if chi2_stat > 0:
                n = observed.sum().sum()
                cramers_v = np.sqrt(chi2_stat / (n * (min(observed.shape) - 1)))
                st.write(f"• **Cramér's V:** {cramers_v:.4f}")
                if cramers_v < 0.1:
                    interpretation = "Association très faible"
                elif cramers_v < 0.3:
                    interpretation = "Association faible"
                elif cramers_v < 0.5:
                    interpretation = "Association modérée"
                else:
                    interpretation = "Association forte"
                st.write(f"• **Interprétation:** {interpretation}")
                
                resultats['tests'].append({
                    'Test': "Cramér's V",
                    'Statistique': round(cramers_v, 4),
                    'Seuils': '< 0.1: très faible, < 0.3: faible, < 0.5: modérée, ≥ 0.5: forte',
                    'Interprétation': interpretation
                })
            
        except Exception as e:
            st.write(f"• Erreur dans le test Chi2: {str(e)}")
            resultats['tests'].append({
                'Test': 'Erreur Chi2',
                'Erreur': str(e),
                'Interprétation': 'Test impossible'
            })
    
    with col2:
        st.write("📊 **Visualisations:**")
        
        # Affichage des graphiques
        create_stacked_bar_echarts(df_clean, var_cible, var_cat, save_folder, colors)
        create_heatmap_echarts(df_clean, var_cible, var_cat, save_folder, colors)
    
    return resultats

def analyze_numeric_vs_numeric(df_clean, var_cible, var_num, ponderation, save_folder, colors):
    """Analyse de corrélation entre deux variables numériques avec affichage complet et retour de données"""
    
    # Dictionnaire pour stocker les résultats
    resultats = {
        'type_analyse': 'Numérique vs Numérique',
        'n_observations': len(df_clean),
        'tests': [],
        'tableau_principal': None,
        'test_principal': None,
        'p_value_principal': None,
        'significatif_principal': None
    }
    
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.write("📊 **Analyse de corrélation:**")
        
        # Corrélation
        correlation = df_clean[var_cible].corr(df_clean[var_num])
        st.write(f"• **Corrélation de Pearson:** {correlation:.4f}")
        
        # Interprétation de la force
        if abs(correlation) > 0.7:
            interpretation = "Forte"
        elif abs(correlation) > 0.5:
            interpretation = "Modérée"
        elif abs(correlation) > 0.3:
            interpretation = "Faible"
        else:
            interpretation = "Très faible"
        
        st.write(f"• **Intensité:** {interpretation}")
        st.write(f"• **Direction:** {'Positive' if correlation > 0 else 'Négative'}")
        
        # Test de significativité
        try:
            from scipy.stats import pearsonr
            _, p_value = pearsonr(df_clean[var_cible], df_clean[var_num])
            
            # Stockage du test principal
            resultats['test_principal'] = "Corrélation de Pearson"
            resultats['p_value_principal'] = round(p_value, 4)
            resultats['significatif_principal'] = p_value < 0.05
            
            st.write(f"• **p-value:** {p_value:.4f}")
            st.write(f"• **Significativité:** {'Significative' if p_value < 0.05 else 'Non significative'}")
            
            # Coefficient de détermination
            r_squared = correlation ** 2
            st.write(f"• **R² (coefficient de détermination):** {r_squared:.4f}")
            st.write(f"• **Variance expliquée:** {r_squared*100:.1f}%")
            
            # Ajouter aux tests
            resultats['tests'].append({
                'Test': 'Corrélation de Pearson',
                'Coefficient': round(correlation, 4),
                'p-value': round(p_value, 4),
                'R²': round(r_squared, 4),
                'Variance expliquée (%)': round(r_squared*100, 1),
                'Significatif': p_value < 0.05,
                'Intensité': interpretation,
                'Direction': 'Positive' if correlation > 0 else 'Négative',
                'Interprétation': f"Corrélation {interpretation.lower()} {'significative' if p_value < 0.05 else 'non significative'}"
            })
            
        except Exception as e:
            st.write(f"• Erreur dans le test de corrélation: {str(e)}")
            resultats['tests'].append({
                'Test': 'Erreur corrélation',
                'Erreur': str(e)
            })
        
        # Tableau de contingence des quartiles
        st.write("📋 **Répartition par quartiles:**")
        try:
            quartiles_x = pd.qcut(df_clean[var_cible], 4, labels=['Q1', 'Q2', 'Q3', 'Q4'], duplicates='drop')
            quartiles_y = pd.qcut(df_clean[var_num], 4, labels=['Q1', 'Q2', 'Q3', 'Q4'], duplicates='drop')
            quartiles_crosstab = pd.crosstab(quartiles_x, quartiles_y)
            resultats['tableau_principal'] = quartiles_crosstab
            st.dataframe(quartiles_crosstab, use_container_width=True)
        except Exception as e:
            st.write(f"Impossible de créer les quartiles: {str(e)}")
            resultats['tableau_principal'] = pd.DataFrame({'Erreur': [str(e)]})
    
    with col2:
        st.write("📊 **Nuage de points:**")
        
        # Nuage de points avec ECharts
        data = [[float(x), float(y)] for x, y in zip(df_clean[var_cible], df_clean[var_num]) if not (pd.isna(x) or pd.isna(y))]
        
        option = {
            "title": {
                "text": f"📈 {var_cible} vs {var_num}",
                "left": "center",
                "textStyle": {"fontSize": 16, "color": "#2E4057"}
            },
            "tooltip": {"trigger": "item"},
            "xAxis": {
                "type": "value",
                "name": var_cible,
                "axisLabel": {"fontSize": 11}
            },
            "yAxis": {
                "type": "value",
                "name": var_num,
                "axisLabel": {"fontSize": 11}
            },
            "series": [
                {
                    "type": "scatter",
                    "data": data,
                    "itemStyle": {"color": colors[0]},
                    "symbolSize": 8
                }
            ],
            "grid": {
                "left": "15%",
                "right": "10%",
                "bottom": "15%",
                "top": "15%"
            }
        }
        
        st_echarts(options=option, height="400px", key=f"scatter_{var_cible}_{var_num}")
        
        # Export du nuage de points
        if save_folder:
            fig = px.scatter(df_clean, x=var_cible, y=var_num, trendline="ols",
                           title=f"📈 Nuage de points: {var_cible} vs {var_num}")
            fig.update_layout(
                title={'x': 0.5, 'xanchor': 'center', 'font': {'size': 16, 'color': '#2E4057'}},
                height=500, font=dict(family="Arial, sans-serif"),
                paper_bgcolor='rgba(0,0,0,0)', plot_bgcolor='rgba(0,0,0,0)'
            )
            
            clean_cible = "".join(c for c in var_cible if c.isalnum() or c in (' ', '-', '_')).rstrip()
            clean_num = "".join(c for c in var_num if c.isalnum() or c in (' ', '-', '_')).rstrip()
            clean_cible = clean_cible.replace(' ', '_')
            clean_num = clean_num.replace(' ', '_')
            save_path = f"{save_folder}/scatter_{clean_cible}_vs_{clean_num}.png"
            
            try:
                pio.write_image(fig, save_path, width=800, height=500)
            except Exception as e:
                st.warning(f"Impossible de sauvegarder le graphique: {str(e)}")
    
    return resultats

# Interface principale
def main():
    # Créer les dossiers de sortie
    create_output_folders()
    
    # Titre principal
    st.markdown('<div class="main-header">📊 Analyseur Statistique Pro</div>', unsafe_allow_html=True)
    
    # Sidebar pour le chargement des données
    with st.sidebar:
        st.header("🔧 Configuration")
        
        # Upload de fichier
        uploaded_file = st.file_uploader(
            "📁 Charger vos données",
            type=['csv', 'xlsx', 'xls'],
            help="Formats supportés: CSV, Excel (.xlsx, .xls)"
        )
        
        # Variables globales pour stocker les données d'analyse
        if 'analyses_data' not in st.session_state:
            st.session_state.analyses_data = {
                'numeric_vars': [],
                'categorical_vars': [],
                'resultats_bivariees': None,
                'var_cible': None,
                'df': None,
                'ponderation': None,
                'palette_name': None,
                'selected_sheet': None
            }
        
        # Variables par défaut
        df = None
        selected_sheet = None
        ponderation = None
        palette_name = "Défaut (Verte/Orange)"
        selected_colors = COLOR_PALETTES[palette_name]
        show_missing = True
        show_correlation = True
        numeric_vars = []
        categorical_vars = []
        
        if uploaded_file is not None:
            # Gestion spéciale pour les fichiers Excel avec plusieurs feuilles
            if uploaded_file.name.endswith(('.xlsx', '.xls')):
                # Récupérer la liste des feuilles
                excel_sheets = get_excel_sheets(uploaded_file)
                
                if len(excel_sheets) > 1:
                    st.subheader("📋 Sélection de la feuille Excel")
                    st.info(f"📊 Ce fichier Excel contient **{len(excel_sheets)} feuilles**. Sélectionnez celle qui contient vos données à analyser.")
                    
                    selected_sheet = st.selectbox(
                        "Choisissez la feuille à analyser:",
                        excel_sheets,
                        help="Chaque feuille peut contenir des données différentes"
                    )
                    
                    # Aperçu des feuilles disponibles
                    with st.expander("👀 Aperçu des feuilles disponibles"):
                        for i, sheet in enumerate(excel_sheets, 1):
                            try:
                                # Lire quelques lignes pour aperçu
                                preview_df = pd.read_excel(uploaded_file, sheet_name=sheet, nrows=3)
                                st.write(f"**{i}. {sheet}** ({preview_df.shape[0]}+ lignes × {preview_df.shape[1]} colonnes)")
                                if len(preview_df.columns) <= 10:  # Afficher seulement si pas trop de colonnes
                                    st.write(f"Colonnes: {', '.join(preview_df.columns[:10])}")
                                else:
                                    st.write(f"Colonnes: {', '.join(preview_df.columns[:10])}... (+{len(preview_df.columns)-10} autres)")
                            except Exception as e:
                                st.write(f"**{i}. {sheet}** (Erreur: {str(e)})")
                else:
                    selected_sheet = excel_sheets[0] if excel_sheets else None
            
            # Chargement des données
            if uploaded_file.name.endswith('.csv') or selected_sheet is not None:
                with st.spinner("Chargement des données..."):
                    df = load_data(uploaded_file, selected_sheet)
            
            if df is not None:
                st.success(f"✅ {df.shape[0]} lignes, {df.shape[1]} colonnes")
                if selected_sheet:
                    st.info(f"📋 Feuille chargée: **{selected_sheet}**")
                
                # Définir les types de variables
                numeric_vars = df.select_dtypes(include=[np.number]).columns.tolist()
                categorical_vars = df.select_dtypes(include=['object', 'category']).columns.tolist()
                
                # Mettre à jour les données dans le session_state
                st.session_state.analyses_data['df'] = df
                st.session_state.analyses_data['selected_sheet'] = selected_sheet
                
                # Sélection de la palette de couleurs
                st.subheader("🎨 Palette de couleurs")
                
                # Sélecteur de palette
                palette_name = st.selectbox(
                    "Choisissez votre palette:",
                    list(COLOR_PALETTES.keys()),
                    index=0,
                    help="La palette sélectionnée sera utilisée pour tous les graphiques"
                )
                
                # Aperçu de la palette sélectionnée
                selected_colors = COLOR_PALETTES[palette_name]
                show_palette_preview("Aperçu", selected_colors)
                
                # Mettre à jour la palette dans le session_state
                st.session_state.analyses_data['palette_name'] = palette_name
                
                # Sélection de la pondération
                st.subheader("⚖️ Pondération (optionnel)")
                ponderation = st.selectbox(
                    "Variable de pondération:",
                    ['Aucune'] + numeric_vars,
                    help="Sélectionnez une variable numérique pour pondérer les analyses"
                )
                ponderation = None if ponderation == 'Aucune' else ponderation
                
                # Retirer la variable de pondération des variables à analyser
                if ponderation and ponderation in numeric_vars:
                    numeric_vars.remove(ponderation)
                
                # Mettre à jour les données dans le session_state
                st.session_state.analyses_data['numeric_vars'] = numeric_vars
                st.session_state.analyses_data['categorical_vars'] = categorical_vars
                st.session_state.analyses_data['ponderation'] = ponderation
                
                # Options d'affichage
                st.subheader("🎛️ Options d'affichage")
                show_missing = st.checkbox("Afficher l'analyse des données manquantes", value=True)
                show_correlation = st.checkbox("Afficher la matrice de corrélation", value=True)
    
    # Contenu principal
    if uploaded_file is not None and df is not None:
        # Récupération de la palette sélectionnée
        selected_colors = COLOR_PALETTES[palette_name]
        
        # Onglets principaux
        tab1, tab2, tab3, tab4 = st.tabs(["📋 Aperçu des données", "📊 Analyses univariées", "🔗 Analyses bivariées", "📥 Exports"])
        
        with tab1:
            st.markdown('<div class="sub-header">📋 Aperçu des données</div>', unsafe_allow_html=True)
            
            # Informations générales
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                st.metric("📏 Lignes", df.shape[0])
            with col2:
                st.metric("📊 Colonnes", df.shape[1])
            with col3:
                memory_mb = df.memory_usage(deep=True).sum() / 1024**2
                st.metric("💾 Mémoire", f"{memory_mb:.1f} MB")
            with col4:
                missing_pct = (df.isnull().sum().sum() / (df.shape[0] * df.shape[1])) * 100
                st.metric("❌ Manquantes", f"{missing_pct:.1f}%")
            
            # Aperçu du dataset
            st.subheader("👀 Aperçu du dataset")
            st.dataframe(df.head(10), use_container_width=True)
            
            # Types de variables
            col1, col2 = st.columns(2)
            with col1:
                st.subheader("🔢 Variables numériques")
                st.write(f"**{len(numeric_vars)} variables:** {', '.join(numeric_vars) if numeric_vars else 'Aucune'}")
            
            with col2:
                st.subheader("📝 Variables qualitatives")
                st.write(f"**{len(categorical_vars)} variables:** {', '.join(categorical_vars) if categorical_vars else 'Aucune'}")
            
            # Analyse des données manquantes
            if show_missing:
                st.subheader("🚨 Analyse des données manquantes")
                missing_data = df.isnull().sum()
                missing_pct = (missing_data / len(df)) * 100
                missing_df = pd.DataFrame({
                    'Variable': missing_data.index,
                    'Valeurs manquantes': missing_data.values,
                    'Pourcentage (%)': missing_pct.values
                }).query('`Valeurs manquantes` > 0').sort_values('Valeurs manquantes', ascending=False)
                
                if len(missing_df) > 0:
                    st.dataframe(missing_df, use_container_width=True)
                    
                    # Graphique des données manquantes
                    if len(missing_df) <= 20:  # Limiter pour la lisibilité
                        value_counts = missing_df.set_index('Variable')['Pourcentage (%)']
                        create_bar_chart_echarts(value_counts, "Données manquantes (%)", None, selected_colors, 0)
                else:
                    st.success("✅ Aucune donnée manquante détectée!")
            
            # Matrice de corrélation
            if show_correlation and len(numeric_vars) > 1:
                st.subheader("🔗 Matrice de corrélation")
                corr_matrix = df[numeric_vars].corr()
                
                # Graphique avec ECharts
                data = []
                for i, row_name in enumerate(corr_matrix.index):
                    for j, col_name in enumerate(corr_matrix.columns):
                        data.append([j, i, round(corr_matrix.loc[row_name, col_name], 3)])
                
                option = {
                    "title": {
                        "text": "🔗 Matrice de corrélation",
                        "left": "center",
                        "textStyle": {"fontSize": 16}
                    },
                    "tooltip": {"position": "top"},
                    "grid": {"height": "50%", "top": "15%"},
                    "xAxis": {
                        "type": "category",
                        "data": list(corr_matrix.columns),
                        "splitArea": {"show": True}
                    },
                    "yAxis": {
                        "type": "category",
                        "data": list(corr_matrix.index),
                        "splitArea": {"show": True}
                    },
                    "visualMap": {
                        "min": -1,
                        "max": 1,
                        "calculable": True,
                        "orient": "horizontal",
                        "left": "center",
                        "bottom": "15%",
                        "inRange": {"color": ["#313695", "#f7f7f7", "#a50026"]}
                    },
                    "series": [
                        {
                            "name": "Corrélation",
                            "type": "heatmap",
                            "data": data,
                            "label": {"show": True}
                        }
                    ]
                }
                
                st_echarts(options=option, height="600px", key="correlation_matrix")
        
        with tab2:
            st.markdown('<div class="sub-header">📊 Analyses statistiques univariées</div>', unsafe_allow_html=True)
            
            # Bouton pour lancer l'analyse
            if st.button("🚀 Lancer l'analyse univariée", type="primary", use_container_width=True):
                
                save_folder = "Sorties/Stat_Univariées"
                
                # Analyse des variables numériques
                if numeric_vars:
                    st.subheader("🔢 Variables numériques")
                    
                    with st.spinner("Analyse des variables numériques..."):
                        stats_table = analyze_numeric_variables(df, numeric_vars, ponderation)
                    
                    st.dataframe(stats_table, use_container_width=True)
                    
                    # Export Excel des statistiques
                    excel_buffer = io.BytesIO()
                    stats_table.to_excel(excel_buffer, index=True)
                    st.download_button(
                        label="📥 Télécharger les statistiques (Excel)",
                        data=excel_buffer.getvalue(),
                        file_name=f"stats_numeriques_{'pondere' if ponderation else 'simple'}_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    )
                
                # Analyse des variables qualitatives
                if categorical_vars:
                    st.subheader("📝 Variables qualitatives")
                    
                    for i, var in enumerate(categorical_vars):
                        with st.container():
                            st.markdown(f'<div class="analysis-container">', unsafe_allow_html=True)
                            st.write(f"**📊 Variable: {var}**")
                            
                            # Calcul des fréquences
                            if ponderation is None:
                                value_counts = df[var].value_counts()
                                total = value_counts.sum()
                            else:
                                mask_valid = df[var].notna() & df[ponderation].notna()
                                if mask_valid.sum() > 0:
                                    value_counts = df[mask_valid].groupby(var)[ponderation].sum().sort_values(ascending=False)
                                    total = value_counts.sum()
                                else:
                                    st.warning(f"Aucune donnée valide pour {var}")
                                    continue
                            
                            # Tableau de fréquences
                            freq_table = pd.DataFrame({
                                'Effectif': value_counts,
                                'Fréquence (%)': (value_counts / total * 100).round(2)
                            })
                            
                            col1, col2 = st.columns([1, 2])
                            with col1:
                                st.dataframe(freq_table.head(10))
                                
                                # Informations sur la variable
                                n_categories = len(value_counts)
                                missing_count = df[var].isnull().sum()
                                st.write(f"• **Modalités:** {n_categories}")
                                st.write(f"• **Manquantes:** {missing_count} ({(missing_count/len(df)*100):.1f}%)")
                                if ponderation:
                                    st.write(f"• **Effectif pondéré:** {total:.0f}")
                            
                            with col2:
                                # Logique de choix du graphique selon le nombre de modalités
                                if n_categories <= 3:
                                    # Graphique en secteurs (alternance disque/anneau)
                                    create_pie_chart_echarts(value_counts, var, save_folder, selected_colors, i)
                                else:
                                    # Graphique en barres (alternance vertical/horizontal)
                                    create_bar_chart_echarts(value_counts, var, save_folder, selected_colors, i)
                            
                            st.markdown('</div>', unsafe_allow_html=True)
        
        with tab3:
            st.markdown('<div class="sub-header">🔗 Analyses statistiques bivariées</div>', unsafe_allow_html=True)
            
            # Sélection des variables
            col1, col2 = st.columns(2)
            with col1:
                st.subheader("🎯 Variable cible")
                all_vars = list(df.columns)
                if ponderation and ponderation in all_vars:
                    all_vars.remove(ponderation)
                
                var_cible = st.selectbox(
                    "Sélectionnez la variable cible:",
                    all_vars,
                    help="Variable avec laquelle croiser toutes les autres"
                )
            
            with col2:
                st.subheader("📋 Variables à croiser")
                autres_vars = [col for col in all_vars if col != var_cible]
                variables_selectionnees = st.multiselect(
                    "Sélectionnez les variables (laisser vide pour toutes):",
                    autres_vars,
                    help="Si aucune variable n'est sélectionnée, toutes seront analysées"
                )
                
                if not variables_selectionnees:
                    variables_selectionnees = autres_vars
            
            # Bouton pour lancer l'analyse
            if st.button("🚀 Lancer l'analyse bivariée", type="primary", use_container_width=True):
                if var_cible and variables_selectionnees:
                    
                    save_folder = "Sorties/Stat_Bivariées"
                    
                    # Dictionnaire pour stocker tous les résultats d'analyse
                    resultats_analyses = {}
                    
                    # Mettre à jour la variable cible dans le session_state
                    st.session_state.analyses_data['var_cible'] = var_cible
                    
                    # Progress bar
                    progress_bar = st.progress(0)
                    status_text = st.empty()
                    
                    for i, var in enumerate(variables_selectionnees):
                        status_text.text(f"Analyse de {var_cible} vs {var}...")
                        progress_bar.progress((i + 1) / len(variables_selectionnees))
                        
                        with st.container():
                            st.markdown(f'<div class="analysis-container">', unsafe_allow_html=True)
                            st.markdown(f"### 🔍 {var_cible} ✖️ {var}")
                            
                            # Nettoyage des données
                            cols_to_use = [var_cible, var]
                            if ponderation:
                                cols_to_use.append(ponderation)
                            
                            df_clean = df[cols_to_use].dropna()
                            
                            if len(df_clean) == 0:
                                st.warning(f"Aucune donnée valide pour l'analyse {var_cible} vs {var}")
                                continue
                            
                            is_target_numeric = df[var_cible].dtype in ['int64', 'float64', 'int32', 'float32']
                            is_var_numeric = df[var].dtype in ['int64', 'float64', 'int32', 'float32']
                            
                            if is_var_numeric and not is_target_numeric:
                                # Variable numérique vs catégorielle
                                resultats_analyses[var] = analyze_numeric_vs_categorical(df_clean, var_cible, var, ponderation, save_folder, selected_colors)
                                
                            elif not is_var_numeric and not is_target_numeric:
                                # Deux variables catégorielles
                                resultats_analyses[var] = analyze_categorical_vs_categorical(df_clean, var_cible, var, ponderation, save_folder, selected_colors)
                                
                            elif is_var_numeric and is_target_numeric:
                                # Deux variables numériques - corrélation
                                resultats_analyses[var] = analyze_numeric_vs_numeric(df_clean, var_cible, var, ponderation, save_folder, selected_colors)
                            
                            st.markdown('</div>', unsafe_allow_html=True)
                    
                    status_text.text("✅ Analyse terminée!")
                    progress_bar.empty()
                    status_text.empty()
                    
                    # Stocker les résultats dans le session_state
                    st.session_state.analyses_data['resultats_bivariees'] = resultats_analyses
                    
                    # Boutons d'export
                    if resultats_analyses:
                        st.subheader("📥 Export des résultats")
                        
                        col1, col2 = st.columns(2)
                        
                        with col1:
                            if st.button("📊 Télécharger Excel", type="secondary", use_container_width=True):
                                excel_buffer = create_bivariate_excel_report(resultats_analyses, var_cible, ponderation, palette_name, selected_sheet)
                                st.download_button(
                                    label="📥 Télécharger le rapport Excel",
                                    data=excel_buffer.getvalue(),
                                    file_name=f"analyses_bivariees_{var_cible}_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx",
                                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                                )
                        
                        with col2:
                            if st.button("📄 Télécharger Word", type="secondary", use_container_width=True):
                                with st.spinner("Génération du rapport Word..."):
                                    word_doc = create_word_report(
                                        df, numeric_vars, categorical_vars, 
                                        resultats_analyses, var_cible, 
                                        ponderation, palette_name, selected_sheet
                                    )
                                    
                                    # Sauvegarder le document dans un buffer
                                    word_buffer = io.BytesIO()
                                    word_doc.save(word_buffer)
                                    word_buffer.seek(0)
                                    
                                    st.download_button(
                                        label="📥 Télécharger le rapport Word",
                                        data=word_buffer.getvalue(),
                                        file_name=f"rapport_analyses_{var_cible}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                    )
        
        with tab4:
            st.markdown('<div class="sub-header">📥 Exports et téléchargements</div>', unsafe_allow_html=True)
            
            st.info("💡 Les graphiques sont automatiquement sauvegardés en PNG dans les dossiers Sorties/")
            
            # Affichage de la palette utilisée
            st.subheader("🎨 Palette utilisée")
            show_palette_preview(palette_name, selected_colors)
            
            # Statistiques descriptives
            st.subheader("📊 Exports des analyses")
            
            col1, col2, col3 = st.columns(3)
            
            with col1:
                st.write("**📊 Analyses univariées**")
                
                # Export Excel
                if st.button("📊 Excel univarié", use_container_width=True):
                    with st.spinner("Génération du rapport Excel..."):
                        # Créer un fichier Excel avec plusieurs feuilles
                        excel_buffer = io.BytesIO()
                        with pd.ExcelWriter(excel_buffer, engine='openpyxl') as writer:
                            
                            # Feuille des variables numériques
                            if numeric_vars:
                                stats_table = analyze_numeric_variables(df, numeric_vars, ponderation)
                                stats_table.to_excel(writer, sheet_name='Variables_numeriques', index=True)
                            
                            # Feuille des variables qualitatives
                            if categorical_vars:
                                all_freq_tables = {}
                                for var in categorical_vars:
                                    if ponderation is None:
                                        value_counts = df[var].value_counts()
                                        total = value_counts.sum()
                                    else:
                                        mask_valid = df[var].notna() & df[ponderation].notna()
                                        if mask_valid.sum() > 0:
                                            value_counts = df[mask_valid].groupby(var)[ponderation].sum().sort_values(ascending=False)
                                            total = value_counts.sum()
                                        else:
                                            continue
                                    
                                    freq_table = pd.DataFrame({
                                        'Effectif': value_counts,
                                        'Fréquence (%)': (value_counts / total * 100).round(2)
                                    })
                                    all_freq_tables[var] = freq_table
                                
                                # Écrire chaque tableau sur une ligne différente
                                start_row = 0
                                for var_name, freq_table in all_freq_tables.items():
                                    freq_table.to_excel(writer, sheet_name='Variables_qualitatives', 
                                                      startrow=start_row, startcol=0)
                                    # Ajouter le nom de la variable
                                    worksheet = writer.sheets['Variables_qualitatives']
                                    worksheet.cell(row=start_row+1, column=1, value=f"Variable: {var_name}")
                                    start_row += len(freq_table) + 3
                            
                            # Feuille d'informations générales
                            info_df = pd.DataFrame({
                                'Métrique': ['Nombre de lignes', 'Nombre de colonnes', 'Variables numériques', 
                                           'Variables qualitatives', 'Feuille Excel', 'Pondération utilisée', 'Palette de couleurs', 'Date d\'analyse'],
                                'Valeur': [df.shape[0], df.shape[1], len(numeric_vars), len(categorical_vars),
                                         selected_sheet if selected_sheet else 'N/A (CSV)', 
                                         ponderation if ponderation else 'Aucune', palette_name, datetime.now().strftime('%Y-%m-%d %H:%M')]
                            })
                            info_df.to_excel(writer, sheet_name='Informations', index=False)
                        
                        st.download_button(
                            label="📥 Télécharger Excel univarié",
                            data=excel_buffer.getvalue(),
                            file_name=f"rapport_univarie_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            key="download_excel_univarie"
                        )
                
                # Export Word
                if st.button("📄 Word univarié", use_container_width=True):
                    with st.spinner("Génération du rapport Word..."):
                        word_doc = create_word_report(
                            df, numeric_vars, categorical_vars,
                            None, None,  # Pas d'analyses bivariées
                            ponderation, palette_name, selected_sheet
                        )
                        
                        # Sauvegarder le document dans un buffer
                        word_buffer = io.BytesIO()
                        word_doc.save(word_buffer)
                        word_buffer.seek(0)
                        
                        st.download_button(
                            label="📥 Télécharger Word univarié",
                            data=word_buffer.getvalue(),
                            file_name=f"rapport_univarie_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="download_word_univarie"
                        )
            
            with col2:
                st.write("**🔗 Analyses bivariées**")
                
                # Vérifier si des analyses bivariées ont été effectuées
                data = st.session_state.analyses_data
                if data.get('resultats_bivariees') and data.get('var_cible'):
                    st.success("✅ Analyses bivariées disponibles")
                    
                    # Export Excel bivariées
                    if st.button("📊 Excel bivariées", use_container_width=True):
                        excel_buffer = create_bivariate_excel_report(
                            data['resultats_bivariees'], 
                            data['var_cible'], 
                            ponderation, 
                            palette_name, 
                            selected_sheet
                        )
                        st.download_button(
                            label="📥 Télécharger Excel bivariées",
                            data=excel_buffer.getvalue(),
                            file_name=f"analyses_bivariees_{data['var_cible']}_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            key="download_excel_bivariees"
                        )
                    
                    # Export Word bivariées
                    if st.button("📄 Word bivariées", use_container_width=True):
                        with st.spinner("Génération du rapport Word bivariées..."):
                            word_doc = create_word_report(
                                df, numeric_vars, categorical_vars,
                                data['resultats_bivariees'], 
                                data['var_cible'],
                                ponderation, palette_name, selected_sheet
                            )
                            
                            # Sauvegarder le document dans un buffer
                            word_buffer = io.BytesIO()
                            word_doc.save(word_buffer)
                            word_buffer.seek(0)
                            
                            st.download_button(
                                label="📥 Télécharger Word bivariées",
                                data=word_buffer.getvalue(),
                                file_name=f"rapport_bivariees_{data['var_cible']}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key="download_word_bivariees"
                            )
                    
                    st.write("**Contenu des rapports bivariées :**")
                    st.write("• Synthèse de toutes les analyses")
                    st.write("• Tableaux croisés et statistiques")
                    st.write("• Tests statistiques complets")
                    st.write("• Graphiques intégrés")
                    st.write("• Interprétations détaillées")
                else:
                    st.info("💡 Effectuez d'abord des analyses bivariées dans l'onglet correspondant.")
                    st.write("**Une fois les analyses effectuées :**")
                    st.write("• Exports Excel et Word disponibles")
                    st.write("• Tableaux croisés et statistiques")
                    st.write("• Tous les tests statistiques")
                    st.write("• Graphiques intégrés")
            
            with col3:
                st.write("**📄 Rapport complet**")
                
                # Export Excel complet
                if st.button("📊 Excel complet", use_container_width=True):
                    with st.spinner("Génération du rapport Excel complet..."):
                        # Créer un fichier Excel avec toutes les analyses
                        excel_buffer = io.BytesIO()
                        with pd.ExcelWriter(excel_buffer, engine='openpyxl') as writer:
                            
                            # Informations générales
                            info_df = pd.DataFrame({
                                'Métrique': ['Nombre de lignes', 'Nombre de colonnes', 'Variables numériques', 
                                           'Variables qualitatives', 'Feuille Excel', 'Pondération utilisée', 
                                           'Palette de couleurs', 'Analyses bivariées', 'Date d\'analyse'],
                                'Valeur': [df.shape[0], df.shape[1], len(numeric_vars), len(categorical_vars),
                                         selected_sheet if selected_sheet else 'N/A (CSV)', 
                                         ponderation if ponderation else 'Aucune', palette_name,
                                         'Oui' if data.get('resultats_bivariees') else 'Non',
                                         datetime.now().strftime('%Y-%m-%d %H:%M')]
                            })
                            info_df.to_excel(writer, sheet_name='00_Informations', index=False)
                            
                            # Variables numériques
                            if numeric_vars:
                                stats_table = analyze_numeric_variables(df, numeric_vars, ponderation)
                                stats_table.to_excel(writer, sheet_name='01_Var_numeriques', index=True)
                            
                            # Variables qualitatives
                            if categorical_vars:
                                start_row = 0
                                for var in categorical_vars:
                                    if ponderation is None:
                                        value_counts = df[var].value_counts()
                                        total = value_counts.sum()
                                    else:
                                        mask_valid = df[var].notna() & df[ponderation].notna()
                                        if mask_valid.sum() > 0:
                                            value_counts = df[mask_valid].groupby(var)[ponderation].sum().sort_values(ascending=False)
                                            total = value_counts.sum()
                                        else:
                                            continue
                                    
                                    freq_table = pd.DataFrame({
                                        'Effectif': value_counts,
                                        'Fréquence (%)': (value_counts / total * 100).round(2)
                                    })
                                    
                                    freq_table.to_excel(writer, sheet_name='02_Var_qualitatives', 
                                                      startrow=start_row, startcol=0)
                                    worksheet = writer.sheets['02_Var_qualitatives']
                                    worksheet.cell(row=start_row+1, column=1, value=f"Variable: {var}")
                                    start_row += len(freq_table) + 3
                            
                            # Analyses bivariées si disponibles
                            if data.get('resultats_bivariees'):
                                # Synthèse bivariées
                                synthese_data = []
                                for var_name, resultats in data['resultats_bivariees'].items():
                                    synthese_data.append({
                                        'Variable': var_name,
                                        'Type d\'analyse': resultats['type_analyse'],
                                        'Observations': resultats['n_observations'],
                                        'Test principal': resultats.get('test_principal', 'N/A'),
                                        'P-value': resultats.get('p_value_principal', 'N/A'),
                                        'Significatif': resultats.get('significatif_principal', 'N/A')
                                    })
                                
                                synthese_df = pd.DataFrame(synthese_data)
                                synthese_df.to_excel(writer, sheet_name='03_Synthese_bivariees', index=False)
                        
                        st.download_button(
                            label="📥 Télécharger Excel complet",
                            data=excel_buffer.getvalue(),
                            file_name=f"rapport_complet_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            key="download_excel_complet"
                        )
                
                # Export Word complet
                if st.button("📄 Word complet", use_container_width=True):
                    with st.spinner("Génération du rapport Word complet..."):
                        word_doc = create_word_report(
                            df, numeric_vars, categorical_vars,
                            data.get('resultats_bivariees'), 
                            data.get('var_cible'),
                            ponderation, palette_name, selected_sheet
                        )
                        
                        # Sauvegarder le document dans un buffer
                        word_buffer = io.BytesIO()
                        word_doc.save(word_buffer)
                        word_buffer.seek(0)
                        
                        st.download_button(
                            label="📥 Télécharger Word complet",
                            data=word_buffer.getvalue(),
                            file_name=f"rapport_complet_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="download_word_complet"
                        )
                
                st.write("**Contenu du rapport complet :**")
                st.write("• Informations générales")
                st.write("• Analyses univariées complètes")
                st.write("• Tableaux de fréquences")
                st.write("• Graphiques intégrés")
                st.write("• Analyses bivariées (si effectuées)")
                st.write("• Tests statistiques")
                st.write("• Conclusion et notes techniques")
            
            # Données filtrées
            st.subheader("📋 Export des données")
            if st.button("📥 Télécharger les données (CSV)"):
                csv_buffer = io.StringIO()
                df.to_csv(csv_buffer, index=False)
                st.download_button(
                    label="📥 Télécharger CSV",
                    data=csv_buffer.getvalue(),
                    file_name=f"donnees_exportees_{datetime.now().strftime('%Y%m%d_%H%M')}.csv",
                    mime="text/csv"
                )
            
            # Informations sur les exports
            st.subheader("📁 Localisation des fichiers")
            st.info("""
            **Les fichiers sont automatiquement sauvegardés dans :**
            - 📊 `Sorties/Stat_Univariées/` : Graphiques des analyses univariées (PNG)
            - 🔗 `Sorties/Stat_Bivariées/` : Graphiques des analyses bivariées (PNG)
            - 📥 `Sorties/Exports/` : Rapports Excel et exports de données
            """)
            
            st.subheader("💡 Note d'installation")
            st.info("""
            **Pour utiliser la génération de rapports Word, installez :**
            ```bash
            pip install python-docx
            ```
            Cette dépendance est nécessaire pour créer les documents Word avec tableaux et images.
            """)
    
    else:
        # Page d'accueil si aucun fichier n'est chargé ou aucune feuille sélectionnée
        if uploaded_file is not None and df is None:
            st.warning("⚠️ Veuillez sélectionner une feuille Excel pour continuer.")
        else:
            st.markdown("""
            <div class="info-box">
                <h3>🚀 Bienvenue dans l'Analyseur Statistique Pro!</h3>
                <p>Cet outil vous permet d'effectuer facilement des analyses statistiques complètes sur vos données.</p>
                
                <h4>📋 Fonctionnalités:</h4>
                <ul>
                    <li>📊 <strong>Analyses univariées</strong> - Statistiques descriptives et graphiques</li>
                    <li>🔗 <strong>Analyses bivariées</strong> - Tableaux croisés et tests statistiques</li>
                    <li>⚖️ <strong>Analyses pondérées</strong> - Support de la pondération</li>
                    <li>🎨 <strong>Palettes personnalisables</strong> - 7 palettes de couleurs au choix</li>
                    <li>📈 <strong>Graphiques interactifs</strong> - Visualisations avec streamlit-echarts</li>
                    <li>📥 <strong>Exports automatiques</strong> - Excel, CSV, images PNG</li>
                    <li>📄 <strong>Rapports Word</strong> - Documents professionnels avec tableaux et graphiques</li>
                    <li>📋 <strong>Support multi-feuilles</strong> - Sélection de feuilles Excel</li>
                </ul>
                
                <h4>🎨 Nouvelle logique des graphiques:</h4>
                <ul>
                    <li><strong>≤ 3 modalités:</strong> Graphiques en secteurs (alternance disque/anneau)</li>
                    <li><strong>> 3 modalités:</strong> Graphiques en barres (alternance vertical/horizontal)</li>
                </ul>
                
                <h4>🔧 Pour commencer:</h4>
                <ol>
                    <li>Chargez votre fichier de données (CSV ou Excel) dans la barre latérale</li>
                    <li>Si c'est un fichier Excel, sélectionnez la feuille à analyser</li>
                    <li>Choisissez votre palette de couleurs préférée</li>
                    <li>Configurez éventuellement une variable de pondération</li>
                    <li>Explorez vos données dans l'onglet "Aperçu"</li>
                    <li>Lancez vos analyses dans les onglets dédiés</li>
                </ol>
                
                <h4>📋 Formats supportés:</h4>
                <ul>
                    <li><strong>Entrée:</strong> CSV (détection automatique du séparateur), Excel (.xlsx, .xls) multi-feuilles</li>
                    <li><strong>Sortie:</strong> Excel (.xlsx), Word (.docx), Images PNG, CSV</li>
                </ul>
            </div>
            """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()